# Объединение парсеров в один.

# TODO: .doc

import collections
from io import StringIO
import json
import os
from pathlib import Path
import re
import xml
import xml.etree
import xml.etree.ElementTree

from bs4 import BeautifulSoup as bs
import docx  # python-docx
import docx2txt2
from docx_parser import DocumentParser  # docx_parser
from markdown import Markdown
from striprtf.striprtf import rtf_to_text  # striprtf
import pptx2txt2
from pydocx import PyDocX

from pdf_parser import extract_text_from_pdf, clean_text


collections.Hashable = collections.abc.Hashable


class MarkdownToPlain:
    def __init__(self):
        Markdown.output_formats["plain"] = MarkdownToPlain.unmark_element
        self.md = Markdown(output_format="plain")
        self.md.stripTopLevelTags = False

    @staticmethod
    def unmark_element(element: xml.etree.ElementTree.Element, stream: StringIO | None = None) -> str:
        if stream is None:
            stream = StringIO()
        if element.text:
            stream.write(element.text)
        for sub in element:
            MarkdownToPlain.unmark_element(sub, stream)
        if element.tail:
            stream.write(element.tail)
        return stream.getvalue()

    def convert(self, source: str) -> str:
        return self.md.convert(source)


md = MarkdownToPlain()


def extract_text_from_markdown(path: os.PathLike) -> str:
    with open(path) as f:
        return md.convert(f.read())


def extract_text_from_rtf(path: os.PathLike) -> str:
    with open(path) as f:
        text = rtf_to_text(f.read())
        text = re.sub(r"\|", " ", text)
        return text


def extract_text_from_docx(path: os.PathLike) -> str:
    """
    Извлекает текст из файла формата DOCX (MS OFFICE 2007+).
    Зачастую бросает ошибку, связанную со структурой документа.

    Parameters:
    - path (os.PathLike): Путь к файлу.

    Returns:
    - str: текст элемента
    """
    res: list[str] = []
    doc = DocumentParser(path)
    for _type, item in doc.parse():
        if "text" in item:
            res.append(item["text"])
        elif "data" in item:
            res.extend(" ".join(row) for row in item["data"])

    return "\n".join(res)


def extract_text_from_docx_2(path: os.PathLike) -> str:
    document = docx.Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    tables = "\n".join(" ".join(c.text() for c in row.cells)
                       for table in document.tables for row in table.rows)

    return f"{text}\n{tables}"


def extract_text_from_doc(path: os.PathLike) -> str:
    html = PyDocX.to_html(path)
    return extract_text_with_bs(html)


def extract_text_with_bs(source: str) -> str:
    soup = bs(source)
    [s.extract() for s in soup(["style", "script"])]
    tmpText = soup.get_text(separator="\n")
    return tmpText


def read_any_doc(path: os.PathLike) -> str:
    text: str = ""
    extention: str = path.suffix

    if extention == ".txt":
        with open(path) as f:
            text = f.read()
    elif extention == ".md":
        text = extract_text_from_markdown(path)
    elif extention == ".pdf":
        text = extract_text_from_pdf(path)
    elif extention == ".rtf":
        text = extract_text_from_rtf(path)
    elif extention == ".pptx" or extention == ".odp":
        text = pptx2txt2.extract_text(path)
    elif extention == ".odt":
        text = docx2txt2.extract_text(path)
    elif extention == ".docx":
        try:
            text = extract_text_from_docx(path)
        except Exception as e:
            try:
                text = extract_text_from_docx_2(path)
            except Exception as e:
                # могут присутствовать артефакты
                text = docx2txt2.extract_text(path)
    elif extention == ".doc":
        text = extract_text_from_doc(path)

        if text == '':
            text = docx2txt2.extract_text(path)
    else:
        with open(path, errors="ignore") as f:
            text = extract_text_with_bs(f.read())
    return clean_text(text)


if __name__ == "__main__":
    import time
    import os

    # Чтение данных из файла JSON

    new_data = []  # Список для хранения успешно обработанных данных
    bad_data = []  # Список для хранения данных с ошибками
    so_so_data = []

    # Итерация по элементам данных
    start_time = time.time()
    for filename in os.listdir("./data/"):
        iteration_start_time = time.time()

        item = {
            "id": filename.split(".")[0],
            "file_name": filename
        }
        try:

            # Извлечение текста из PDF файла
            item["text"] = read_any_doc(Path("data", item.get("file_name")))

            # Если текст успешно извлечен
            if item["text"] != "":
                with open(Path("data", f'{item.get("id")}.txt'), 'w') as f:
                    f.write(item['text'])
                # Добавляем элемент в список успешных данных
                new_data.append(item)
            else:

                so_so_data.append(item)
        except Exception as e:
            # Добавляем информацию об ошибке в элемент данных
            item["error"] = str(e)
            # Добавляем элемент в список данных с ошибками
            bad_data.append(item)

    # Запись успешно обработанных данных в файл JSON
    with open("data_text_true_recovered.json", "w", encoding="utf-8") as file:
        json.dump(new_data, file, ensure_ascii=False, indent=1)

    # Запись данных с ошибками в файл JSON
    with open("data_text_false_recovered.json", "w", encoding="utf-8") as file:
        json.dump(bad_data, file, ensure_ascii=False, indent=1)

    with open("data_text_so_so_recovered.json", "w", encoding="utf-8") as file:
        json.dump(so_so_data, file, ensure_ascii=False, indent=1)
